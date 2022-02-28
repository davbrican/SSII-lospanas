import csv
import cv2
import numpy as np
from fpdf import FPDF
from docx import Document
from docx.shared import Inches

filesNumber = int(int(input("Número de archivos (mínimo 5): ")) / 5)

for i in range(filesNumber):
    with open("./files/csv"+str(i)+".csv", "w", newline='') as f:
        writer = csv.writer(f, delimiter=',', quoting=csv.QUOTE_MINIMAL)
        writer.writerow(["file", "yepa"])
        writer.writerow([str(i),"quepasagente"])
        
    with open("./files/text"+str(i)+".txt", "w") as f:
        f.write("fileNumber"+str(i))
        

    height = 720
    width = 1080

    blank_image = np.zeros((height,width,3), np.uint8)
    blank_image[:,i] = (255,0,0)      # (B, G, R)
    blank_image[:,i:width] = (0,255,0)

    cv2.imwrite("./files/image"+str(i)+".png", blank_image)
    
      
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size = 15)
    
    pdf.cell(200, 10, txt = "File Number "+str(i), 
            ln = 1, align = 'C')
    
    # add another cell
    pdf.cell(200, 10, txt = "This is a PDF test",
            ln = 2, align = 'C')
    
    # save the pdf with name .pdf
    pdf.output("./files/pdf"+str(i)+".pdf")  
    

    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('This is the docx number '+str(i)+' ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )


    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('./files/docs'+str(i)+'.docx')
